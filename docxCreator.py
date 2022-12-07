from docx import Document
from docx.shared import Inches
guestnames = ['Игорь', 'Илья', 'Мади', 'Искандер']
document = Document()
for x in guestnames:
    document.add_heading('Новосибирский зоопарк', 0)
    p = document.add_paragraph(x + ', посетите')
    p.add_run(' один из лучших зоопарков').bold = True
    p.add_run(', стран СНГ.')
    document.add_paragraph('Только у нас вы сможете увидеть 750 видов животных, 350 из которых занесены в международную Красную книгу.')
    document.add_paragraph('')
    document.add_paragraph('Наш зоопарк также занимается разведением кошачьих и куньих, поэтому здесь одна из лучших в мире коллекций представителей этих семейств.')
    document.add_paragraph('')
    document.add_paragraph('Не ждите момента, а создавайте его!')
    document.add_paragraph('Мы ждем вас по адресу ул. Тимирязева, д.71/1')
    document.add_picture('zoo.jpg', width=Inches(5.25))
    document.add_page_break()
document.save('visit.docx')